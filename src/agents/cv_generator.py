"""
CV Generator Agent
Produces tailored resume outputs in DOCX format using professional templates
"""

from typing import Dict, Any, Optional
import logging
import json
from pathlib import Path
from datetime import datetime

from .base_agent import BaseAgent

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not installed. Install with: pip install python-docx")

logger = logging.getLogger(__name__)


class CVGeneratorAgent(BaseAgent):
    """
    ADK Agent for CV generation

    Responsibilities:
    - Apply job-specific tailoring to CV content
    - Select appropriate CV template based on job position
    - Populate template with tailored CV data
    - Generate professional DOCX output compatible with Word/Google Docs
    - Optimize keyword matching for ATS systems
    """

    # Template mapping based on job position/title keywords
    TEMPLATE_MAPPING = {
        "executive": ["executive", "director", "vp", "vice president", "ceo", "cto", "cfo", "chief"],
        "engineering": ["engineer", "developer", "programmer", "software", "backend", "frontend", "fullstack", "devops", "sre"],
        "management": ["manager", "lead", "head of", "supervisor", "coordinator"],
        "design": ["designer", "ux", "ui", "creative", "artist", "graphic"],
        "data": ["data scientist", "data analyst", "data engineer", "ml engineer", "machine learning"],
        "marketing": ["marketing", "growth", "seo", "content", "brand", "digital marketing"],
        "sales": ["sales", "account executive", "business development", "account manager"],
        "finance": ["accountant", "financial analyst", "finance", "auditor", "controller"],
        "operations": ["operations", "logistics", "supply chain", "operations manager"],
        "hr": ["hr", "human resources", "recruiter", "talent acquisition"],
        "consulting": ["consultant", "advisor", "analyst", "strategist"],
    }

    def __init__(self, llm_provider=None, storage_backend=None, config: Optional[Dict[str, Any]] = None):
        super().__init__(
            name="cv_generator",
            description="Generates tailored CVs in DOCX format optimized for ATS and specific job positions",
            llm_provider=llm_provider,
            storage_backend=storage_backend,
            config=config
        )

        self.output_dir = Path(config.get("output_dir", "./data/outputs"))
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.template_dir = Path(config.get("template_dir", "./templates/cv_templates"))
        self.template_dir.mkdir(parents=True, exist_ok=True)

    async def generate(
        self,
        cv_data: Dict[str, Any],
        job_requirements: Optional[Dict[str, Any]] = None,
        gap_analysis: Optional[Dict[str, Any]] = None,
        user_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate tailored CV in DOCX format

        This is a key A2A action called by the orchestrator

        Args:
            cv_data: JSON Resume formatted CV
            job_requirements: Job requirements from job analysis (includes position)
            gap_analysis: Gap analysis results
            user_id: User identifier

        Returns:
            Dictionary with generated file paths (docx and json)
        """
        try:
            logger.info(f"[{self.name}] Generating CV for user: {user_id}")

            # Step 1: Tailor CV content for job
            tailored_cv = await self._tailor_cv(
                cv_data,
                job_requirements,
                gap_analysis
            )

            # Step 2: Select appropriate template based on job position
            template_type = self._select_template(job_requirements)
            logger.info(f"[{self.name}] Selected template: {template_type}")

            # Step 3: Generate DOCX output using template
            output_files = {}

            if DOCX_AVAILABLE:
                output_files['docx'] = await self._generate_docx(
                    tailored_cv,
                    user_id,
                    template_type,
                    job_requirements
                )
            else:
                logger.error(f"[{self.name}] python-docx not available. Cannot generate DOCX files.")
                raise ImportError("python-docx is required. Install with: pip install python-docx")

            # Step 4: Also save JSON for reference
            output_files['json'] = await self._generate_json(
                tailored_cv,
                user_id
            )

            # Store outputs if storage backend available
            if self.storage_backend:
                for format_name, file_path in output_files.items():
                    remote_path = f"outputs/{user_id}/{Path(file_path).name}"
                    self.storage_backend.upload_file(
                        local_path=file_path,
                        remote_path=remote_path
                    )

            result = {
                "output_files": output_files,
                "tailored_cv": tailored_cv,
                "template_used": template_type,
                "formats_generated": list(output_files.keys()),
                "user_id": user_id
            }

            logger.info(f"[{self.name}] Generated CV successfully (template: {template_type})")

            return result

        except Exception as e:
            logger.error(f"[{self.name}] CV generation failed: {e}", exc_info=True)
            raise

    def _select_template(self, job_requirements: Optional[Dict[str, Any]]) -> str:
        """
        Select appropriate template based on job position

        Args:
            job_requirements: Job requirements including position/title

        Returns:
            Template type identifier
        """
        if not job_requirements:
            return "professional"  # Default template

        # Extract job title/position
        job_title = job_requirements.get("title", "").lower()
        job_position = job_requirements.get("position", "").lower()
        job_role = job_requirements.get("role", "").lower()

        search_text = f"{job_title} {job_position} {job_role}".lower()

        # Match against template mapping
        for template_name, keywords in self.TEMPLATE_MAPPING.items():
            for keyword in keywords:
                if keyword in search_text:
                    logger.info(f"[{self.name}] Matched '{keyword}' -> template: {template_name}")
                    return template_name

        # Default to professional template
        logger.info(f"[{self.name}] No specific match, using default 'professional' template")
        return "professional"

    async def _tailor_cv(
        self,
        cv_data: Dict[str, Any],
        job_requirements: Optional[Dict[str, Any]],
        gap_analysis: Optional[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """
        Tailor CV content for specific job

        Args:
            cv_data: Original CV data
            job_requirements: Job requirements
            gap_analysis: Gap analysis results

        Returns:
            Tailored CV data
        """
        if not self.llm_provider or not job_requirements:
            logger.warning(f"[{self.name}] No LLM or job requirements, returning original CV")
            return cv_data

        # Use LLM to tailor CV with strict format enforcement matching profile JSON structure
        prompt = f"""Tailor this CV for the specific job requirements while maintaining STRICT JSON format that matches our profile storage format.

CRITICAL FORMAT REQUIREMENTS - You MUST follow this EXACT structure:

{{
  "basics": {{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "+XX XXXXXXXXX",
    "url": "website (optional)",
    "location": {{
      "address": "Full address string"
    }},
    "profiles": [
      {{"network": "LinkedIn", "url": "https://..."}},
      {{"network": "GitHub", "url": "https://..."}}
    ]
  }},
  "work": [
    {{
      "company": "Company Name",
      "position": "Job Title",
      "startDate": "YYYY-MM",
      "endDate": "YYYY-MM" or null for current,
      "highlights": ["Achievement 1", "Achievement 2"]
    }}
  ],
  "education": [
    {{
      "institution": "University Name",
      "area": "Field of Study",
      "studyType": "Degree Type",
      "startDate": "YYYY-MM",
      "endDate": "YYYY-MM",
      "gpa": "X.XX/10" (optional)
    }}
  ],
  "skills": [
    {{"name": "Category Name", "keywords": ["skill1", "skill2", "skill3"]}},
    {{"name": "Another Category", "keywords": ["skill4", "skill5"]}}
  ],
  "projects": [
    {{"name": "Project Name", "description": ["description text", "highlight 1", "highlight 2"]}}
  ],
  "certificates": [
    {{"name": "Certificate Name", "details": ["Issuer: Organization", "Date: YYYY-MM"]}}
  ]
}}

FIELD NAME RULES (DO NOT DEVIATE):
- Work: Use "company", "position", "startDate", "endDate", "highlights"
- Education: Use "institution", "area", "studyType", "startDate", "endDate", "gpa"
- Skills: Use "name" and "keywords" (NOT "category", "items", "level", etc.)
- Projects: Use "name" and "description" array (NOT "highlights", "summary", "technologies")
- Certificates: Use "name" and "details" array (NOT "issuer", "date" as separate fields)
- Basics/Location: Use "address" (single string, NOT "city", "country", etc.)
- Basics/Profiles: Use "network" and "url"

Original CV:
{json.dumps(cv_data, indent=2)[:2500]}

Job Requirements:
{json.dumps(job_requirements, indent=2)[:1000]}

Gap Analysis Recommendations:
{json.dumps(gap_analysis.get('recommendations', []) if gap_analysis else [], indent=2)[:500]}

TAILORING INSTRUCTIONS:
1. Highlight relevant work experience by reordering "highlights" arrays (most relevant first)
2. Add job-specific keywords naturally in "highlights" arrays
3. Emphasize relevant skills by reordering skill categories
4. Keep ALL field names exactly as shown in the format above
5. DO NOT fabricate experience, skills, or achievements
6. DO NOT add new fields or rename existing fields
7. DO NOT change date formats (keep YYYY-MM)
8. DO NOT split location.address into separate fields

Return ONLY valid JSON matching the exact format above."""

        try:
            tailored_cv = await self.llm_provider.complete_json(
                prompt=prompt,
                temperature=0.3,
                max_tokens=3000
            )

            logger.info(f"[{self.name}] CV tailored for job requirements")

            # Validate complete format in tailored CV (matching profile JSON structure)
            format_valid = True

            # Validate top-level structure
            required_sections = ["basics", "work", "education", "skills", "projects", "certificates"]
            for section in required_sections:
                if section not in tailored_cv:
                    logger.error(f"[{self.name}] Missing required section: {section}")
                    format_valid = False

            # Validate basics structure
            basics = tailored_cv.get("basics", {})
            if "location" in basics and "address" not in basics["location"]:
                logger.error(f"[{self.name}] basics.location missing 'address' field")
                format_valid = False

            # Validate work entries
            work = tailored_cv.get("work", [])
            logger.info(f"[{self.name}] Tailored CV has {len(work)} work experiences")
            for idx, job in enumerate(work[:2], 1):  # Check first 2
                required_work_fields = ["company", "position", "startDate", "highlights"]
                missing = [f for f in required_work_fields if f not in job]
                if missing:
                    logger.error(f"[{self.name}] Work {idx} missing fields: {missing}")
                    format_valid = False
                else:
                    logger.info(f"[{self.name}]   Work {idx}: {job['position']} at {job['company']}")

            # Validate skills format
            skills = tailored_cv.get("skills", [])
            logger.info(f"[{self.name}] Tailored CV has {len(skills)} skill groups")
            for idx, skill in enumerate(skills[:3], 1):  # Check first 3
                if not isinstance(skill, dict):
                    logger.error(f"[{self.name}] Skill {idx} is not a dict: {type(skill)}")
                    format_valid = False
                elif "name" not in skill or "keywords" not in skill:
                    logger.error(f"[{self.name}] Skill {idx} missing 'name' or 'keywords': {skill.keys()}")
                    format_valid = False
                else:
                    logger.info(f"[{self.name}]   Skill {idx}: {skill['name']} ({len(skill['keywords'])} keywords)")

            if not format_valid:
                logger.warning(f"[{self.name}] Tailored CV has invalid format, using original CV")
                return cv_data

            logger.info(f"[{self.name}] Tailored CV format validation passed")
            return tailored_cv

        except Exception as e:
            logger.error(f"[{self.name}] CV tailoring failed, using original: {e}", exc_info=True)
            return cv_data

    async def _generate_docx(
        self,
        cv_data: Dict[str, Any],
        user_id: Optional[str],
        template_type: str,
        job_requirements: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Generate DOCX format CV using professional template

        Args:
            cv_data: CV data (tailored)
            user_id: User identifier
            template_type: Template type to use
            job_requirements: Job requirements (for context)

        Returns:
            Path to generated .docx file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        # Create new document
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        basics = cv_data.get("basics", {})

        # === HEADER SECTION ===
        # Name
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(basics.get('name', 'Resume'))
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(31, 78, 120)  # Professional blue
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Professional Title / Label
        if basics.get("label"):
            label_paragraph = doc.add_paragraph()
            label_run = label_paragraph.add_run(basics['label'])
            label_run.font.size = Pt(12)
            label_run.font.italic = True
            label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact Information
        contact_parts = []
        if basics.get("email"):
            contact_parts.append(basics['email'])
        if basics.get("phone"):
            contact_parts.append(basics['phone'])
        if basics.get("url"):
            contact_parts.append(basics['url'])
        if basics.get("location", {}).get("city"):
            location = basics['location']
            loc_str = location.get('city', '')
            if location.get('region'):
                loc_str += f", {location['region']}"
            contact_parts.append(loc_str)

        if contact_parts:
            contact_paragraph = doc.add_paragraph(" | ".join(contact_parts))
            contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_paragraph.runs:
                run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing

        # === PROFESSIONAL SUMMARY ===
        if basics.get("summary"):
            summary_heading = doc.add_heading('Professional Summary', level=2)
            summary_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)

            summary_p = doc.add_paragraph(basics['summary'])
            summary_p.runs[0].font.size = Pt(11)

        # === WORK EXPERIENCE ===
        work = cv_data.get("work", [])
        if work:
            work_heading = doc.add_heading('Work Experience', level=2)
            work_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)

            for job in work:
                # Job title and company
                position = job.get("position", "")
                company = job.get("company") or job.get("name", "")

                job_title_p = doc.add_paragraph()
                job_title_run = job_title_p.add_run(f"{position}")
                job_title_run.font.size = Pt(12)
                job_title_run.font.bold = True

                # Company and dates
                start_date = job.get("startDate", "")
                end_date = job.get("endDate") or "Present"

                company_p = doc.add_paragraph()
                company_run = company_p.add_run(f"{company}")
                company_run.font.italic = True
                company_run.font.size = Pt(11)

                if start_date or end_date:
                    dates_run = company_p.add_run(f" | {start_date} - {end_date}")
                    dates_run.font.size = Pt(10)
                    dates_run.font.color.rgb = RGBColor(100, 100, 100)

                # Job summary
                if job.get("summary"):
                    summary_p = doc.add_paragraph(job["summary"])
                    summary_p.runs[0].font.size = Pt(10)

                # Highlights/achievements
                highlights = job.get("highlights", [])
                if highlights:
                    for highlight in highlights:
                        bullet_p = doc.add_paragraph(highlight, style='List Bullet')
                        bullet_p.runs[0].font.size = Pt(10)

                doc.add_paragraph()  # Spacing between jobs

        # === EDUCATION ===
        education = cv_data.get("education", [])
        if education:
            edu_heading = doc.add_heading('Education', level=2)
            edu_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)

            for edu in education:
                institution = edu.get("institution", "")
                study_type = edu.get("studyType") or edu.get("degree", "")
                area = edu.get("area") or edu.get("field", "")
                start_date = edu.get("startDate", "")
                end_date = edu.get("endDate") or "Present"

                # Degree and field
                degree_p = doc.add_paragraph()
                degree_run = degree_p.add_run(f"{study_type} in {area}" if area else study_type)
                degree_run.font.size = Pt(11)
                degree_run.font.bold = True

                # Institution and dates
                inst_p = doc.add_paragraph()
                inst_run = inst_p.add_run(institution)
                inst_run.font.italic = True
                inst_run.font.size = Pt(10)

                if start_date or end_date:
                    dates_run = inst_p.add_run(f" | {start_date} - {end_date}")
                    dates_run.font.size = Pt(9)
                    dates_run.font.color.rgb = RGBColor(100, 100, 100)

                doc.add_paragraph()  # Spacing

        # === SKILLS ===
        skills = cv_data.get("skills", [])
        if skills:
            skills_heading = doc.add_heading('Skills', level=2)
            skills_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)

            for skill_group in skills:
                # Debug: Log the skill_group to identify problematic data
                logger.info(f"[{self.name}] Processing skill_group: {skill_group} (type: {type(skill_group)})")

                # Handle case where skill_group might be a string instead of dict
                if isinstance(skill_group, str):
                    logger.warning(f"[{self.name}] Skill is a string instead of dict: '{skill_group}' - skipping")
                    continue

                if not isinstance(skill_group, dict):
                    logger.warning(f"[{self.name}] Skill is not a dict: {type(skill_group)} - skipping")
                    continue

                logger.info(f"[{self.name}] Dict skill_group contents: {dict(skill_group)}")

                name = skill_group.get("name", "")
                keywords = skill_group.get("keywords", [])

                logger.info(f"[{self.name}] Extracted name='{name}', keywords={keywords}")

                # Check if name exists and keywords is a non-empty list
                if name and len(keywords) > 0:
                    logger.info(f"[{self.name}] Adding skill to DOCX: {name}")
                    skill_p = doc.add_paragraph()
                    name_run = skill_p.add_run(f"{name}: ")
                    name_run.font.bold = True
                    name_run.font.size = Pt(10)

                    keywords_run = skill_p.add_run(", ".join(keywords))
                    keywords_run.font.size = Pt(10)
                else:
                    logger.warning(f"[{self.name}] Skipping skill - name='{name}', keywords_count={len(keywords)}")

        # === ADDITIONAL SECTIONS ===
        # Projects
        projects = cv_data.get("projects", [])
        if projects:
            projects_heading = doc.add_heading('Projects', level=2)
            projects_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)

            for project in projects:
                project_name = project.get("name", "")
                project_desc = project.get("description", "")

                if project_name:
                    proj_p = doc.add_paragraph()
                    proj_run = proj_p.add_run(project_name)
                    proj_run.font.bold = True
                    proj_run.font.size = Pt(10)

                if project_desc:
                    desc_p = doc.add_paragraph(project_desc)
                    desc_p.runs[0].font.size = Pt(9)

        # Certificates
        certificates = cv_data.get("certificates", [])
        if certificates:
            cert_heading = doc.add_heading('Certifications', level=2)
            cert_heading.runs[0].font.color.rgb = RGBColor(31, 78, 120)

            for cert in certificates:
                cert_name = cert.get("name", "")
                cert_issuer = cert.get("issuer", "")
                cert_date = cert.get("date", "")

                if cert_name:
                    cert_p = doc.add_paragraph()
                    cert_run = cert_p.add_run(cert_name)
                    cert_run.font.size = Pt(10)
                    cert_run.font.bold = True

                    if cert_issuer:
                        issuer_run = cert_p.add_run(f" - {cert_issuer}")
                        issuer_run.font.italic = True
                        issuer_run.font.size = Pt(9)

                    if cert_date:
                        date_run = cert_p.add_run(f" ({cert_date})")
                        date_run.font.size = Pt(9)
                        date_run.font.color.rgb = RGBColor(100, 100, 100)

        # Generate filename
        name_part = basics.get("name", user_id or "resume").replace(" ", "_")
        job_title = job_requirements.get("title", "position") if job_requirements else "position"
        job_title_part = job_title.replace(" ", "_")[:20]  # Limit length
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename = f"{name_part}_{job_title_part}_{template_type}_{timestamp}.docx"
        filepath = self.output_dir / filename

        # Save document
        doc.save(filepath)

        logger.info(f"[{self.name}] Generated DOCX CV: {filepath}")

        return str(filepath)

    async def _generate_json(
        self,
        cv_data: Dict[str, Any],
        user_id: Optional[str]
    ) -> str:
        """
        Generate JSON Resume format for reference

        Args:
            cv_data: CV data
            user_id: User identifier

        Returns:
            Path to generated file
        """
        name_part = cv_data.get("basics", {}).get("name", user_id or "resume").replace(" ", "_")
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename = f"{name_part}_CV_{timestamp}.json"
        filepath = self.output_dir / filename

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(cv_data, f, indent=2)

        logger.info(f"[{self.name}] Generated JSON CV: {filepath}")

        return str(filepath)

    async def process(self, **kwargs) -> Dict[str, Any]:
        """Main processing entry point"""
        return await self.generate(**kwargs)
